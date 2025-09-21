import os
import io
import tempfile
import logging
import google.generativeai as genai
from django.conf import settings
import pandas as pd
import fitz  # PyMuPDF for PDFs
from docx import Document
import pytesseract
from PIL import Image
import json
from pdf2image import convert_from_path
from .rag_service import validate_term_sheet_with_rag, make_json_serializable
import traceback
import re
import numpy as np
import pdfplumber

# Configure logging
logger = logging.getLogger(__name__)

# Configure Gemini API
genai.configure(api_key=settings.GEMINI_API_KEY)

def extract_text_from_file(uploaded_file):
    """Extract text from an uploaded file."""
    try:
        # Get file extension from filename
        filename = uploaded_file.name if hasattr(uploaded_file, 'name') else 'unknown'
        file_extension = os.path.splitext(filename)[1].lower()
        
        # Log file information
        file_size_mb = uploaded_file.size / (1024 * 1024) if hasattr(uploaded_file, 'size') else 'unknown'
        logger.info(f"Extracting text from {filename} ({file_extension}), size: {file_size_mb} MB")

        # For PDF files
        if file_extension == '.pdf':
            return extract_text_from_pdf(uploaded_file)
            
        # For DOCX files
        elif file_extension == '.docx':
            return extract_text_from_docx(uploaded_file)
            
        # For plain text files
        elif file_extension in ['.txt', '.md', '.csv', '.json']:
            try:
                text = uploaded_file.read().decode('utf-8')
                logger.info(f"Successfully extracted {len(text)} characters from text file")
                return text
            except UnicodeDecodeError:
                logger.warning("UTF-8 decoding failed, trying with another encoding")
                uploaded_file.seek(0)
                text = uploaded_file.read().decode('latin-1')
                logger.info(f"Successfully extracted {len(text)} characters using latin-1 encoding")
                return text
        
        # For image files - attempt OCR
        elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif']:
            try:
                logger.info("Attempting OCR on image file")
                
                # Save the image to a temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
                    temp_file.write(uploaded_file.read())
                    temp_path = temp_file.name
                
                try:
                    # Use Pytesseract for OCR
                    image = Image.open(temp_path)
                    text = pytesseract.image_to_string(image)
                    
                    # Clean up temporary file
                    os.unlink(temp_path)
                    
                    if not text.strip():
                        logger.warning("OCR returned empty text")
                        return "No text could be extracted from the image."
                    
                    logger.info(f"Successfully extracted {len(text)} characters via OCR")
                    return text
                except Exception as ocr_error:
                    # Clean up temporary file
                    os.unlink(temp_path)
                    logger.error(f"OCR failed: {str(ocr_error)}")
                    raise
            except Exception as img_error:
                logger.error(f"Image processing failed: {str(img_error)}")
                return f"Error extracting text from image: {str(img_error)}"
                
        # For unknown file types
        else:
            logger.warning(f"Unsupported file type: {file_extension}")
            return f"Unsupported file type: {file_extension}. Please upload a PDF, DOCX, or TXT file."
            
    except Exception as e:
        logger.error(f"Error in extract_text_from_file: {str(e)}")
        logger.error(traceback.format_exc())
        return f"Error extracting text: {str(e)}"

def extract_text_from_pdf(pdf_file):
    """Extract text from a PDF file using PyMuPDF if available, falling back to PDFPlumber."""
    try:
        # Try using PyMuPDF (fitz) first
        try:
            logger.info("Attempting to extract PDF text using PyMuPDF")
            
            # Save to temporary file since fitz needs a file path
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(pdf_file.read())
                temp_path = temp_file.name
            
            text = ""
            try:
                # Open the PDF with PyMuPDF
                doc = fitz.open(temp_path)
                
                # Extract text from each page
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text += page.get_text() + "\n"
                
                doc.close()
                
                # Clean up temporary file
                os.unlink(temp_path)
                
                if not text.strip():
                    logger.warning("PyMuPDF returned empty text, the PDF might be image-based")
                else:
                    logger.info(f"Successfully extracted {len(text)} characters using PyMuPDF")
                    return text
            except Exception as fitz_error:
                # Clean up temporary file
                os.unlink(temp_path)
                logger.warning(f"PyMuPDF extraction failed: {str(fitz_error)}")
                # Continue to PDFPlumber
        except (ImportError, NameError):
            logger.warning("PyMuPDF not available, using PDFPlumber instead")
        
        # Try PDFPlumber
        try:
            logger.info("Attempting to extract PDF text using PDFPlumber")
            pdf_file.seek(0)  # Reset file pointer
            
            # Save to temporary file since pdfplumber needs a file path
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(pdf_file.read())
                temp_path = temp_file.name
            
            text = ""
            with pdfplumber.open(temp_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
            
            # Clean up temporary file
            os.unlink(temp_path)
            
            if not text.strip():
                logger.warning("PDFPlumber returned empty text, the PDF might be image-based")
                # Try OCR as a last resort
                return extract_pdf_with_ocr(pdf_file)
            else:
                logger.info(f"Successfully extracted {len(text)} characters using PDFPlumber")
                return text
                
        except Exception as plumber_error:
            logger.warning(f"PDFPlumber extraction failed: {str(plumber_error)}")
            # Try OCR as a last resort
            return extract_pdf_with_ocr(pdf_file)
            
    except Exception as e:
        logger.error(f"Error in extract_text_from_pdf: {str(e)}")
        logger.error(traceback.format_exc())
        return f"Error extracting text from PDF: {str(e)}"

def extract_pdf_with_ocr(pdf_file):
    """Extract text from a PDF using OCR when text extraction fails."""
    try:
        logger.info("Attempting OCR on PDF file")
        pdf_file.seek(0)  # Reset file pointer
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_file.write(pdf_file.read())
            temp_path = temp_file.name
        
        try:
            # Convert PDF to images
            text = ""
            pdf_document = convert_from_path(temp_path, 300)
            
            for page_num, page_image in enumerate(pdf_document):
                logger.info(f"OCR processing PDF page {page_num+1}")
                # Save page as temporary image
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as img_temp:
                    page_image.save(img_temp.name, 'PNG')
                    img_path = img_temp.name
                
                # Perform OCR on the image
                try:
                    page_text = pytesseract.image_to_string(Image.open(img_path))
                    text += page_text + "\n"
                finally:
                    # Clean up temporary image file
                    os.unlink(img_path)
            
            # Clean up temporary PDF file
            os.unlink(temp_path)
            
            if not text.strip():
                logger.warning("OCR returned empty text")
                return "No text could be extracted from the PDF."
            
            logger.info(f"Successfully extracted {len(text)} characters via OCR")
            return text
        except Exception as ocr_error:
            # Clean up temporary file
            os.unlink(temp_path)
            logger.error(f"PDF OCR failed: {str(ocr_error)}")
            raise
            
    except Exception as e:
        logger.error(f"Error in extract_pdf_with_ocr: {str(e)}")
        logger.error(traceback.format_exc())
        return f"Error extracting text via OCR: {str(e)}"

def extract_text_from_docx(docx_file):
    """Extract text from a DOCX file."""
    try:
        logger.info("Extracting text from DOCX file")
        
        # Save to temporary file since python-docx needs a file path
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(docx_file.read())
            temp_path = temp_file.name
        
        try:
            # Open and extract text from docx
            doc = Document(temp_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text:
                    full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            # Clean up temporary file
            os.unlink(temp_path)
            
            text = "\n".join(full_text)
            if not text.strip():
                logger.warning("DOCX extraction returned empty text")
                return "No text could be extracted from the document."
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX file")
            return text
            
        except Exception as docx_error:
            # Clean up temporary file
            os.unlink(temp_path)
            logger.error(f"DOCX extraction failed: {str(docx_error)}")
            raise
            
    except Exception as e:
        logger.error(f"Error in extract_text_from_docx: {str(e)}")
        logger.error(traceback.format_exc())
        return f"Error extracting text from DOCX: {str(e)}"

def process_extracted_text(text):
    """Process extracted text to structure it in JSON format"""
    try:
        logger.info(f"Processing extracted text ({len(text)} chars)")
        if len(text) < 50:
            logger.warning(f"Extracted text is very short, may lead to poor results: '{text}'")
        
        # Create prompt for structuring the data
        prompt = f"""
        You are a financial document analyzer specialized in trading term sheets. Your task is to accurately extract key information from the term sheet text provided.
        
        IMPORTANT INSTRUCTIONS:
        1. Read the entire document carefully before answering.
        2. Extract the exact values as they appear in the document.
        3. If a field is not explicitly mentioned, use "Not specified" as the value.
        4. Look for both direct mentions and implied information.
        5. Focus on numerical values and financial details.
        6. Be precise with financial terms and amounts.
        7. ALWAYS extract the Trade ID. It might be found under terms like "Trade Identifier", "Trade Reference", "Deal ID", or "Transaction ID".
        8. For numeric values, maintain all decimal places as shown in the document (e.g., extract 1.2000 as "1.2000" not "1.2").
        
        Here is the term sheet text:
        ```
        {text}
        ```
        
        Extract ONLY these specific trading fields:
        1. trade_id: Unique identifier for the trade (CRITICAL - always look for this)
        2. trade_date: Date when the trade was executed
        3. reference_spot_price: Reference price of the underlying asset at trade execution
        4. notional_amount: Principal amount of the trade
        5. strike_price: Price at which the option can be exercised
        6. option_type: Indicates if the option is a Call or Put
        7. position_type: Indicates if the position is Buying or Selling
        8. expiry_date: Date when the option expires
        9. business_calendar: Calendar used for business day calculations (e.g., NYSE, LSE)
        10. delivery_date: Date for delivery of the asset
        11. premium_rate: Rate paid for the option
        12. transaction_currency: Primary currency used in the transaction (Transaction CCY)
        13. counter_currency: Secondary currency in the transaction (Counter CCY)
        14. underlying_currency: Currency pair or identifier of the underlying asset
        
        Format your response as a clean JSON object with these field names as keys. Do not include any explanations or markdown formatting in your response, just the JSON object.
        """
        
        # Create Gemini model instance
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        # Generate structured data from text
        logger.info("Sending prompt to Gemini API")
        response = model.generate_content(prompt)
        
        # Extract text from response
        response_text = response.text
        logger.info(f"Received response from Gemini API ({len(response_text)} chars)")
        
        # Variable to hold parsing errors for debugging
        parsing_error = ""
        structured_data = {}
        
        # Try different strategies to parse the JSON
        # Strategy 1: Direct JSON parsing
        try:
            structured_data = json.loads(response_text)
            logger.info("Successfully parsed JSON directly")
        except json.JSONDecodeError as e:
            parsing_error = f"Direct parsing: {str(e)}"
            logger.info(f"Direct JSON parsing failed: {str(e)}")
        
        # Strategy 2: Clean and parse JSON
        if not structured_data:
            try:
                # Remove markdown code blocks if present
                cleaned_text = response_text.replace("```json", "").replace("```", "").strip()
                structured_data = json.loads(cleaned_text)
                logger.info("Successfully parsed JSON after cleaning markdown")
            except json.JSONDecodeError as e:
                parsing_error = f"{parsing_error}; Markdown cleaning: {str(e)}"
                logger.info(f"JSON parsing after markdown cleaning failed: {str(e)}")
        
        # Strategy 3: Extra cleaning for common issues and parse
        if not structured_data:
            try:
                # Fix common JSON issues: replace single quotes, fix True/False/None
                extra_cleaned = response_text.replace("'", '"')
                extra_cleaned = re.sub(r'\bTrue\b', 'true', extra_cleaned)
                extra_cleaned = re.sub(r'\bFalse\b', 'false', extra_cleaned)
                extra_cleaned = re.sub(r'\bNone\b', 'null', extra_cleaned)
                # Remove markdown code blocks if present
                extra_cleaned = extra_cleaned.replace("```json", "").replace("```", "").strip()
                structured_data = json.loads(extra_cleaned)
                logger.info("Successfully parsed JSON after extra cleaning")
            except json.JSONDecodeError as e:
                parsing_error = f"{parsing_error}; Extra cleaning: {str(e)}"
                logger.info(f"JSON parsing after extra cleaning failed: {str(e)}")
        
        # Strategy 4: Find first { and last } and extract JSON
        if not structured_data:
            try:
                start_idx = response_text.find('{')
                end_idx = response_text.rfind('}')
                if start_idx != -1 and end_idx != -1 and start_idx < end_idx:
                    json_text = response_text[start_idx:end_idx+1]
                    # Replace JavaScript special values with JSON-compatible values
                    json_text = re.sub(r'\bNaN\b', '"Not specified"', json_text)
                    json_text = re.sub(r'\bnull\b', '"Not specified"', json_text)
                    json_text = re.sub(r'\bundefined\b', '"Not specified"', json_text)
                    structured_data = json.loads(json_text)
                    logger.info("Successfully extracted and parsed JSON using character positions")
            except (json.JSONDecodeError, AttributeError) as e:
                parsing_error = f"{parsing_error}; Character position parsing: {str(e)}"
                logger.info(f"Character position JSON parsing failed: {str(e)}")
        
        # Strategy 5: Use a regex pattern to extract fields if all parsing strategies fail
        if not structured_data:
            structured_data = {}
            logger.warning("All JSON parsing strategies failed, attempting field extraction through regex")
            
            # Try to extract individual fields from response text
            try:
                field_patterns = [
                    (r'"trade_id"\s*:\s*"([^"]*)"', "trade_id"),
                    (r'"trade_date"\s*:\s*"([^"]*)"', "trade_date"),
                    (r'"reference_spot_price"\s*:\s*"([^"]*)"', "reference_spot_price"),
                    (r'"notional_amount"\s*:\s*"([^"]*)"', "notional_amount"),
                    (r'"strike_price"\s*:\s*"([^"]*)"', "strike_price"),
                    (r'"option_type"\s*:\s*"([^"]*)"', "option_type"),
                    (r'"position_type"\s*:\s*"([^"]*)"', "position_type"),
                    (r'"expiry_date"\s*:\s*"([^"]*)"', "expiry_date"),
                    (r'"business_calendar"\s*:\s*"([^"]*)"', "business_calendar"),
                    (r'"delivery_date"\s*:\s*"([^"]*)"', "delivery_date"),
                    (r'"premium_rate"\s*:\s*"([^"]*)"', "premium_rate"),
                    (r'"transaction_currency"\s*:\s*"([^"]*)"', "transaction_currency"),
                    (r'"counter_currency"\s*:\s*"([^"]*)"', "counter_currency"),
                    (r'"underlying_currency"\s*:\s*"([^"]*)"', "underlying_currency")
                ]
                
                for pattern, field_name in field_patterns:
                    match = re.search(pattern, response_text)
                    if match:
                        field_value = match.group(1)
                        if field_value and field_value.lower() not in ["null", "undefined", ""]:
                            structured_data[field_name] = field_value
                            logger.info(f"Extracted individual field {field_name} = {field_value}")
            except Exception as extraction_err:
                logger.error(f"Error extracting individual fields: {str(extraction_err)}")
                
        # Additional extraction for Trade ID if not found
        if "trade_id" not in structured_data or structured_data["trade_id"] == "Not specified":
            # Try to find Trade ID with alternative patterns
            try:
                trade_id_patterns = [
                    r'(?:Trade\s*ID|Trade\s*Identifier|Transaction\s*ID|Deal\s*ID|Trade\s*Ref|Reference\s*Number|Trade\s*Number)[:\s]*([A-Za-z0-9_\-]+)',
                    r'ID[:\s]*([A-Za-z0-9_\-]+)',
                    r'Reference[:\s]*([A-Za-z0-9_\-]+)'
                ]
                
                for pattern in trade_id_patterns:
                    match = re.search(pattern, text, re.IGNORECASE)
                    if match:
                        trade_id = match.group(1).strip()
                        structured_data["trade_id"] = trade_id
                        logger.info(f"Found Trade ID using alternative pattern: {trade_id}")
                        break
            except Exception as e:
                logger.error(f"Error in alternative Trade ID extraction: {str(e)}")
        
        # Clean up extracted values
        for field, value in structured_data.items():
            if isinstance(value, str):
                # Remove invisible characters and extra whitespace
                cleaned_value = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', value)  # Remove control chars
                cleaned_value = re.sub(r'\s+', ' ', cleaned_value)  # Normalize whitespace
                cleaned_value = cleaned_value.strip()  # Trim
                
                # Check if numeric for number fields
                if field in ["reference_spot_price", "notional_amount", "strike_price", "premium_rate"]:
                    try:
                        # Try to normalize numeric values while preserving original decimal places
                        if re.match(r'^[\d,.\s]+%?$', cleaned_value):
                            # Remove commas and spaces
                            numeric_value = cleaned_value.replace(',', '').replace(' ', '')
                            # Handle percentage
                            if numeric_value.endswith('%'):
                                numeric_value = numeric_value[:-1]
                            # Convert to float and back to string to ensure valid format
                            # but preserve original decimal places
                            decimal_places = len(numeric_value.split('.')[-1]) if '.' in numeric_value else 0
                            numeric_value = str(round(float(numeric_value), decimal_places))
                            # Add back % if it was there
                            if cleaned_value.endswith('%'):
                                numeric_value += '%'
                            cleaned_value = numeric_value
                    except (ValueError, TypeError):
                        # If conversion fails, keep original cleaned value
                        pass
                
                structured_data[field] = cleaned_value
        
        # Ensure all required fields are present with default "Not specified" value if missing
        required_fields = [
            "trade_id", "trade_date", "reference_spot_price", "notional_amount", "strike_price",
            "option_type", "position_type", "expiry_date", "business_calendar", "delivery_date", 
            "premium_rate", "transaction_currency", "counter_currency", "underlying_currency"
        ]
        
        for field in required_fields:
            if field not in structured_data or not structured_data[field] or structured_data[field] in ["null", "undefined"]:
                structured_data[field] = "Not specified"
                
        # Check if we have any meaningful data or only default values
        meaningful_data = False
        for field in ["trade_id", "trade_date", "reference_spot_price", "notional_amount", "strike_price"]:
            if structured_data.get(field) != "Not specified":
                meaningful_data = True
                break
                
        if not meaningful_data and "error" not in structured_data:
            structured_data["error"] = "Failed to extract any meaningful data from the document"
            logger.warning("No meaningful data was extracted from the document")
        
        logger.info("Returning structured data")
        return structured_data
            
    except Exception as e:
        logger.error(f"Error in processing extracted text: {e}")
        logger.error(traceback.format_exc())
        # Return a default structure in case of error
        return {
            "error": f"Error processing text: {str(e)}",
            "trade_id": "Not specified",
            "trade_date": "Not specified",
            "reference_spot_price": "Not specified",
            "notional_amount": "Not specified",
            "strike_price": "Not specified",
            "option_type": "Not specified",
            "position_type": "Not specified",
            "expiry_date": "Not specified",
            "business_calendar": "Not specified",
            "delivery_date": "Not specified",
            "premium_rate": "Not specified",
            "transaction_currency": "Not specified",
            "counter_currency": "Not specified",
            "underlying_currency": "Not specified"
        }

def validate_term_sheet(structured_data):
    """
    Validate a structured term sheet
    """
    try:
        # Ensure structured_data is valid for processing
        if not structured_data or not isinstance(structured_data, dict):
            return {
                "overall_status": "error",
                "explanation": "Invalid structured data provided for validation",
                "issues": []
            }
        
        # Call the RAG service for validation
        validation_results = validate_term_sheet_with_rag(structured_data)
        
        # Apply the serialization function first
        validation_results = make_json_serializable(validation_results)
        
        # Additional handling for NaN and Infinity values
        def sanitize_floats(obj):
            import math
            if isinstance(obj, dict):
                return {k: sanitize_floats(v) for k, v in obj.items()}
            elif isinstance(obj, list):
                return [sanitize_floats(item) for item in obj]
            elif isinstance(obj, float):
                if math.isnan(obj):
                    return "NaN"
                elif math.isinf(obj):
                    return "Infinity" if obj > 0 else "-Infinity"
                return obj
            else:
                return obj
                
        validation_results = sanitize_floats(validation_results)
        
        # Ensure the validation results are JSON serializable
        try:
            # Roundtrip through JSON to ensure valid structure
            json_string = json.dumps(validation_results)
            return json.loads(json_string)
        except (TypeError, ValueError) as json_err:
            logger.error(f"Validation results not JSON serializable: {json_err}")
            return {
                "overall_status": "uncertain",
                "explanation": "Validation completed but results couldn't be properly serialized",
                "issues": []
            }
    except Exception as e:
        logger.error(f"Error in validate_term_sheet: {str(e)}")
        return {
            "overall_status": "error",
            "explanation": f"Validation error: {str(e)}",
            "issues": []
        } 